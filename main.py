import os
os.environ["CUDA_VISIBLE_DEVICES"] = ""  # Force CPU
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
import streamlit as st
from datetime import datetime
from app.embedding import embed_and_save
from app.rag_chain import get_answer, clear_memory
from app.file_handler import read_file
import pickle

import faiss

import torch
def _noop_to(self, *args, **kwargs):
    return self
torch.nn.Module.to = _noop_to

from sentence_transformers import SentenceTransformer
from docx import Document
import json
import subprocess
import openai

UPLOAD_DIR = "data/uploaded"
INDEX_DIR = "data/index"
UPDATE_DIR = "data/kms_updates"

os.makedirs(UPDATE_DIR, exist_ok=True)

# ✅ Load offline embedding model
local_model_path = r"C:\models\all-MiniLM-L6-v2"
embedding_model = SentenceTransformer(local_model_path)
embedding_model._target_device = 'cpu'  # ✅ Force CPU

st.set_page_config(layout="wide")
st.title("KMS AI Assistant with Memory & Updates")

# ---------------- Sidebar Menu ----------------
st.sidebar.title("Navigation")
menu = st.sidebar.radio(
    "Choose an option:",
    ["Upload & Embed", "Ask Questions", "Update KMS"]
)

# ✅ Button to clear chat memory in sidebar
if st.sidebar.button("Clear Chat Memory"):
    clear_memory()
    if "chat_history_ui" in st.session_state:
        st.session_state.chat_history_ui = []
    st.sidebar.success("Chat memory cleared!")

# ------------------- Page 1: Upload & Embed -------------------
if menu == "Upload & Embed":
    st.subheader("Upload a File to Embed")
    uploaded_file = st.file_uploader("Upload PDF, DOCX, or TXT", type=["pdf", "docx", "txt"])
    description = st.text_input("Optional Description")

    if uploaded_file:
        file_path = os.path.join(UPLOAD_DIR, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.read())

        st.success(f"Saved: {uploaded_file.name}")

        if st.button("Embed File"):
            doc_name = embed_and_save(file_path)
            st.success(f"File embedded and stored as: {doc_name}")

# ------------------- Page 2: Ask Questions -------------------
elif menu == "Ask Questions":
    st.subheader("Chat with a Document")
    existing_files = [f.split(".")[0] for f in os.listdir(INDEX_DIR) if f.endswith(".index")]

    if existing_files:
        selected_doc = st.selectbox("Choose Document", existing_files)

        # Initialize chat history in session state
        if "chat_history_ui" not in st.session_state:
            st.session_state.chat_history_ui = []

        # ---------------- Chat Display ----------------
        st.markdown(
            """
            <style>
            .chat-container { max-height: 400px; overflow-y: auto; padding: 10px; border-radius: 8px;
                              background-color: #1E1E1E; border: 1px solid #444; }
            .user-bubble { background-color: #0d6efd; color: white; padding: 8px 12px; margin: 5px;
                           border-radius: 15px; max-width: 75%; align-self: flex-end; }
            .bot-bubble { background-color: #3a3b3c; color: white; padding: 8px 12px; margin: 5px;
                          border-radius: 15px; max-width: 75%; align-self: flex-start; }
            </style>
            """,
            unsafe_allow_html=True
        )

        chat_html = '<div class="chat-container" style="display:flex;flex-direction:column;">'
        for user_msg, bot_msg in st.session_state.chat_history_ui:
            chat_html += f'<div class="user-bubble" style="align-self:flex-end;">{user_msg}</div>'
            chat_html += f'<div class="bot-bubble" style="align-self:flex-start;">{bot_msg}</div>'
        chat_html += "</div>"

        st.markdown(chat_html, unsafe_allow_html=True)

        # ---------------- Chat Input ----------------
        user_question = st.text_input("Ask a question:")

        if st.button("Get Answer") and user_question:
            answer = get_answer(user_question, selected_doc)

            # Save to chat history
            st.session_state.chat_history_ui.append((user_question, answer))
            st.rerun()

    else:
        st.info("No embedded documents available yet.")

# ------------------- Page 3: Update KMS -------------------
elif menu == "Update KMS":
    st.subheader("Update KMS Knowledge Base")

    existing_files = [f.split(".")[0] for f in os.listdir(INDEX_DIR) if f.endswith(".index")]

    if existing_files:
        selected_project = st.selectbox("Select Project", existing_files)

        categories = [
            "Business Concept",
            "Language Sync - tech,nontech",
            "Core Business logic",
            "Core Technical aspects",
            "Security parameters",
            "Absolute to do",
            "Error handling",
            "Q & A Session"
        ]
        selected_category = st.selectbox("Select Category", categories)
        update_text = st.text_area("Enter your update:")

        if st.button("Submit Update") and update_text:
            # ✅ Save to Word file for human-readable record
            doc_path = os.path.join(UPDATE_DIR, f"{selected_project}.docx")
            if os.path.exists(doc_path):
                doc = Document(doc_path)
            else:
                doc = Document()
                doc.add_heading(f"Updates for {selected_project}", 0)

            doc.add_heading(selected_category, level=1)
            doc.add_paragraph(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}: {update_text}")
            doc.save(doc_path)

            # ✅ Append update to FAISS vector store
            index_path = os.path.join(INDEX_DIR, f"{selected_project}.index")
            chunks_path = os.path.join(INDEX_DIR, f"{selected_project}_chunks.pkl")

            index = faiss.read_index(index_path)
            with open(chunks_path, "rb") as f:
                chunks = pickle.load(f)

            # Add new chunk with category label
            new_chunk = f"[{selected_category}] {update_text}"
            chunks.append(new_chunk)

            embedding = embedding_model.encode([new_chunk])
            index.add(embedding)

            faiss.write_index(index, index_path)
            with open(chunks_path, "wb") as f:
                pickle.dump(chunks, f)

            st.success(f"Update added to {selected_project} under {selected_category}")
    else:
        st.info("No projects available. Please embed a document first.")

